# templates/templates.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _add_heading(doc, text, size=16, bold=True, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def _add_section_heading(doc, text, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.font.size = Pt(12)
    run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para

def _add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def template_minimal(data):
    doc = Document()
    _add_heading(doc, data['name'], size=20)
    doc.add_paragraph(data.get('title', ''))
    doc.add_paragraph(data.get('contact', ''))
    _add_section_heading(doc, "Summary")
    doc.add_paragraph(data.get('summary', ''))
    _add_section_heading(doc, "Experience")
    for exp in data.get('experience', []):
        doc.add_paragraph(exp)
    _add_section_heading(doc, "Education")
    for edu in data.get('education', []):
        doc.add_paragraph(edu)
    _add_section_heading(doc, "Skills")
    for skill in data.get('skills', []):
        doc.add_paragraph(skill)
    _add_section_heading(doc, "Projects")
    for proj in data.get('projects', []):
        doc.add_paragraph(proj)
    # Certificates
    _add_section_heading(doc, "Certifications")
    for cert in data.get('certificates', []):
        doc.add_paragraph(cert)
    # Personal Details (dict)
    _add_section_heading(doc, "Personal Details")
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")
    # Declaration (string)
    _add_section_heading(doc, "Declaration")
    doc.add_paragraph(data.get('declaration', "I hereby declare that the information provided above is true and correct to the best of my knowledge."))
    return doc


def template_corporate(data):
    doc = Document()
    _add_heading(doc, data.get('name', ''), size=22, color=(0, 51, 102))
    doc.add_paragraph(data.get('title', ''))
    doc.add_paragraph(data.get('contact', ''))

    _add_section_heading(doc, "Professional Summary", color=(0, 51, 102))
    doc.add_paragraph(data.get('summary', ''))

    _add_section_heading(doc, "Work History", color=(0, 51, 102))
    _add_bullet_list(doc, data.get('experience', []))

    _add_section_heading(doc, "Academic Background", color=(0, 51, 102))
    _add_bullet_list(doc, data.get('education', []))

    _add_section_heading(doc, "Core Competencies", color=(0, 51, 102))
    _add_bullet_list(doc, data.get('skills', []))

    # Projects section — safely handle lists or strings
    _add_section_heading(doc, "Projects", color=(0, 51, 102))
    projects = data.get('projects', [])
    if isinstance(projects, list):
        _add_bullet_list(doc, projects)
    elif isinstance(projects, str):
        doc.add_paragraph(projects)

    # Certifications
    _add_section_heading(doc, "Certifications", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('certificates', []))

    # Personal Details
    _add_section_heading(doc, "Personal Details", color=(255, 102, 0))
    personal_details = data.get('personal_details', {})
    if isinstance(personal_details, dict):
        for key, value in personal_details.items():
            doc.add_paragraph(f"{key}: {value}")

    # Declaration
    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(
        data.get(
            'declaration',
            "I hereby declare that the information provided above is true and correct to the best of my knowledge."
        )
    )

    return doc


def template_tech_modern(data):
    doc = Document()
    _add_heading(doc, data['name'], size=24, color=(10, 102, 194))
    doc.add_paragraph(data['title'])
    doc.add_paragraph(data['contact'])
    _add_section_heading(doc, "Profile", color=(10, 102, 194))
    doc.add_paragraph(data['summary'])
    _add_section_heading(doc, "Projects & Experience", color=(10, 102, 194))
    _add_bullet_list(doc, data['experience'])
    _add_section_heading(doc, "Education", color=(10, 102, 194))
    _add_bullet_list(doc, data['education'])
    _add_section_heading(doc, "Technical Skills", color=(10, 102, 194))
    _add_bullet_list(doc, data['skills'])
    _add_section_heading(doc, "Certifications", color=(10, 102, 194))
    doc.add_paragraph(data['certifications'])
    # Personal Details
    _add_section_heading(doc, "Personal Details", color=(255, 102, 0))
    personal_details = data.get('personal_details', {})
    for key, value in personal_details.items():
        doc.add_paragraph(f"{key}: {value}")

    # Declaration
    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(
        data.get(
            'declaration',
            "I hereby declare that the information provided above is true and correct to the best of my knowledge."
        )
    )

    return doc

def template_creative(data):
    doc = Document()
    _add_heading(doc, data['name'], size=26, color=(255, 102, 0))
    doc.add_paragraph(data['title'])
    doc.add_paragraph(data['contact'])
    _add_section_heading(doc, "About Me", color=(255, 102, 0))
    doc.add_paragraph(data['summary'])
    _add_section_heading(doc, "Experience Highlights", color=(255, 102, 0))
    _add_bullet_list(doc, data['experience'])
    _add_section_heading(doc, "Education Path", color=(255, 102, 0))
    _add_bullet_list(doc, data['education'])
    _add_section_heading(doc, "Skillset", color=(255, 102, 0))
    _add_bullet_list(doc, data['skills'])
    _add_section_heading(doc, "Certifications", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('certificates', []))

    _add_section_heading(doc, "Personal Details", color=(255, 102, 0))
    personal_details = data.get('personal_details', {})
    for key, value in personal_details.items():
        doc.add_paragraph(f"{key}: {value}")

    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(
        data.get(
            'declaration',
            "I hereby declare that the information provided above is true and correct to the best of my knowledge."
        )
    )

    return doc

def template_infographic(data):
    doc = Document()
    _add_heading(doc, data['name'], size=24, color=(76, 175, 80))
    doc.add_paragraph(data['title'])
    doc.add_paragraph(data['contact'])
    _add_section_heading(doc, "Snapshot", color=(76, 175, 80))
    doc.add_paragraph(data['summary'])
    _add_section_heading(doc, "Key Experiences", color=(76, 175, 80))
    _add_bullet_list(doc, data['experience'])
    _add_section_heading(doc, "Learning", color=(76, 175, 80))
    _add_bullet_list(doc, data['education'])
    _add_section_heading(doc, "Proficiencies", color=(76, 175, 80))
    _add_bullet_list(doc, data['skills'])
    # Certificates
    _add_section_heading(doc, "Certifications", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('certificates', []))

    # Personal 
    _add_section_heading(doc, "Personal Info", color=(255, 102, 0))
    personal_details = data.get('personal_details', {})
    for key, value in personal_details.items():
        doc.add_paragraph(f"{key}: {value}")

    # Declaration
    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(
        data.get(
            'declaration',
            "I hereby declare that the information provided above is true and correct to the best of my knowledge."
        )
    )

    return doc

def template_bordered(data):
    doc = Document()
    _add_heading(doc, data['name'], size=20, color=(0, 0, 0))
    doc.add_paragraph(data['title'])
    doc.add_paragraph(data['contact'])
    _add_section_heading(doc, "Summary")
    doc.add_paragraph(data['summary'])
    _add_section_heading(doc, "Experience")
    _add_bullet_list(doc, data['experience'])
    _add_section_heading(doc, "Education")
    _add_bullet_list(doc, data['education'])
    _add_section_heading(doc, "Skills")
    _add_bullet_list(doc, data['skills'])
    _add_section_heading(doc, "Projects")
    doc.add_paragraph(data['projects'])
    _add_section_heading(doc, "Certifications & Achievements")
    doc.add_paragraph(data['certifications'])
    _add_section_heading(doc, "Hobbies & Interests")
    doc.add_paragraph(data['hobbies'])
    _add_section_heading(doc, "Personal Details")
    doc.add_paragraph(data['personal_details'])
    _add_section_heading(doc, "Declaration")
    doc.add_paragraph(data['declaration'])
    # Adding a border effect by setting the page background color
    section = doc.sections[0]  
    section.start_type = WD_ALIGN_PARAGRAPH.CENTER
    section.background_color = RGBColor(255, 255, 255)  # White  
    # Adding a border effect by setting page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    # Adding a border effect by setting the page width
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Adding a border effect by setting the page orientation
    section.orientation = WD_ALIGN_PARAGRAPH.CENTER
    # Adding a border effect by setting the page size
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    # Adding simple "border effect" with underscores (pseudo-border)
    doc.add_paragraph("_" * 100)
    return doc

def template_side_panel(data):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Title block (centered at top)
    _add_heading(doc, data.get('name', ''), size=22, color=(33, 33, 33))
    para = doc.add_paragraph(data.get('title', ''))
    contact_info = data.get('contact', '')
    if contact_info:
        para.add_run("\n" + contact_info)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create a table for the two-column layout
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.3)  # Side panel width
    table.columns[1].width = Inches(6.0)  # Main content width

    # Skills
    left = table.cell(0, 0)
    p = left.add_paragraph()
    run = p.add_run("SKILLS")
    run.bold = True
    run.font.size = Pt(11)
    for skill in data.get('skills', []):
        left.add_paragraph("• " + skill, style="List Bullet")
    left.add_paragraph("")

    # Hobbies & Interests
    p = left.add_paragraph()
    run = p.add_run("HOBBIES & INTERESTS")
    run.bold = True
    run.font.size = Pt(11)
    left.add_paragraph(data.get('hobbies', ''))
    left.add_paragraph("")

    # Personal Details
    p = left.add_paragraph()
    run = p.add_run("PERSONAL DETAILS")
    run.bold = True
    run.font.size = Pt(11)
    left.add_paragraph(data.get('personal_details', ''))

    # RIGHT COLUMN (Main content)
    right = table.cell(0, 1)
    # Profile Summary
    p = right.add_paragraph()
    run = p.add_run("PROFILE SUMMARY")
    run.bold = True
    run.font.size = Pt(12)
    right.add_paragraph(data.get('summary', ''))
    right.add_paragraph("")

    # Experience
    p = right.add_paragraph()
    run = p.add_run("EXPERIENCE")
    run.bold = True
    run.font.size = Pt(12)
    for exp in data.get('experience', []):
        right.add_paragraph("• " + exp, style="List Bullet")
    right.add_paragraph("")

    # Education
    p = right.add_paragraph()
    run = p.add_run("EDUCATION")
    run.bold = True
    run.font.size = Pt(12)
    for edu in data.get('education', []):
        right.add_paragraph("• " + edu, style="List Bullet")
    right.add_paragraph("")

    # Projects
    p = right.add_paragraph()
    run = p.add_run("PROJECTS")
    run.bold = True
    run.font.size = Pt(12)
    right.add_paragraph(data.get('projects', ''))
    right.add_paragraph("")

    # Certifications & Achievements
    p = right.add_paragraph()
    run = p.add_run("CERTIFICATIONS & ACHIEVEMENTS")
    run.bold = True
    run.font.size = Pt(12)
    right.add_paragraph(data.get('certifications', ''))
    right.add_paragraph("")

    # Declaration
    p = right.add_paragraph()
    run = p.add_run("DECLARATION")
    run.bold = True
    run.font.size = Pt(12)
    right.add_paragraph(data.get('declaration', ''))

    return doc

def template_academic(data):
    doc = Document()
    _add_heading(doc, data['name'], size=18, color=(54, 69, 79))
    doc.add_paragraph(data['title'])
    doc.add_paragraph(data['contact'])
    _add_section_heading(doc, "Research Profile", color=(54, 69, 79))
    doc.add_paragraph(data['summary'])
    _add_section_heading(doc, "Teaching / Research Experience", color=(54, 69, 79))
    _add_bullet_list(doc, data['experience'])
    _add_section_heading(doc, "Education", color=(54, 69, 79))
    _add_bullet_list(doc, data['education'])
    _add_section_heading(doc, "Publications & Skills", color=(54, 69, 79))
    _add_bullet_list(doc, data['skills'])
    _add_section_heading(doc, "Projects")
    doc.add_paragraph(data['projects'])
    _add_section_heading(doc, "Certifications & Achievements")
    doc.add_paragraph(data['certifications'])
    _add_section_heading(doc, "Hobbies & Interests")
    doc.add_paragraph(data['hobbies'])
    _add_section_heading(doc, "Personal Details")
    doc.add_paragraph(data['personal_details'])
    _add_section_heading(doc, "Declaration")
    doc.add_paragraph(data['declaration'])
    return doc

# Template mapping
TEMPLATE_MAP = {
    "Minimal": template_minimal,
    "Corporate": template_corporate,
    "Tech Modern": template_tech_modern,
    "Creative": template_creative,
    "Infographic": template_infographic,
    "Bordered": template_bordered,
    "Side Panel": template_side_panel,
    "Academic": template_academic
}

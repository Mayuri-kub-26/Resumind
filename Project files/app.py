import os
import time
import json
import io
import base64
import logging
import socket

from dotenv import load_dotenv
load_dotenv()

import streamlit as st
from streamlit_lottie import st_lottie
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.section import WD_ORIENT
import requests
from bs4 import BeautifulSoup

# Selenium imports
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# PDF reading
from PyPDF2 import PdfReader

# --- Basic Configuration ---
# logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# Streamlit page config
st.set_page_config(page_title="ResuMind — Smart Resume Builder", layout="wide", initial_sidebar_state="collapsed")

OUTPUT_JSON = os.environ.get("OUTPUT_JSON_PATH", "scraped_profile.json")
DEFAULT_HEADLESS = os.environ.get("HEADLESS", "True").lower() in ("true", "1", "yes")


# ----------------- Lottie Animation Loader -----------------
def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# ----------------- Resume Template Generation -----------------

def _add_heading(doc, text, size=16, bold=True, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def _add_section_heading(doc, text, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.font.size = Pt(12)
    run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para

def _add_bullet_list(doc, items):
    for item in items:
        if item: # Ensure item is not empty
            doc.add_paragraph(item, style="List Bullet")

def template_minimal(data):
    doc = Document()
    _add_heading(doc, data.get('name', 'Your Name'), size=20)
    doc.add_paragraph(data.get('title', ''))
    doc.add_paragraph(data.get('contact', ''))
    
    _add_section_heading(doc, "Summary")
    doc.add_paragraph(data.get('summary', ''))

    _add_section_heading(doc, "Experience")
    _add_bullet_list(doc, data.get('experience', []))

    _add_section_heading(doc, "Education")
    _add_bullet_list(doc, data.get('education', []))

    _add_section_heading(doc, "Skills")
    _add_bullet_list(doc, data.get('skills', []))

    _add_section_heading(doc, "Projects")
    _add_bullet_list(doc, data.get('projects', []))

    _add_section_heading(doc, "Certifications")
    _add_bullet_list(doc, data.get('certificates', []))

    _add_section_heading(doc, "Personal Details")
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")

    _add_section_heading(doc, "Declaration")
    doc.add_paragraph(data.get('declaration', "I hereby declare that the information provided is true to the best of my knowledge."))
    return doc

def template_corporate(data):
    doc = Document()
    _add_heading(doc, data.get('name', ''), size=22, color=(0, 51, 102))
    doc.add_paragraph(data.get('title', ''))
    doc.add_paragraph(data.get('contact', ''))

    _add_section_heading(doc, "Professional Summary", color=(0, 51, 102))
    doc.add_paragraph(data.get('summary', ''))

    _add_section_heading(doc, "Work History", color=(0, 51, 102))
    _add_bullet_list(doc, data.get('experience', []))

    _add_section_heading(doc, "Academic Background", color=(0, 51, 102))
    _add_bullet_list(doc, data.get('education', []))

    _add_section_heading(doc, "Core Competencies", color=(0, 51, 102))
    _add_bullet_list(doc, data.get('skills', []))

    _add_section_heading(doc, "Projects", color=(0, 51, 102))
    _add_bullet_list(doc, data.get('projects', []))

    _add_section_heading(doc, "Certifications", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('certificates', []))

    _add_section_heading(doc, "Personal Details", color=(255, 102, 0))
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")

    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(data.get('declaration', "I hereby declare that the information provided is true to the best of my knowledge."))
    return doc
    
def template_tech_modern(data):
    doc = Document()
    _add_heading(doc, data.get('name',''), size=24, color=(10, 102, 194))
    doc.add_paragraph(data.get('title',''))
    doc.add_paragraph(data.get('contact',''))
    
    _add_section_heading(doc, "Profile", color=(10, 102, 194))
    doc.add_paragraph(data.get('summary',''))
    
    _add_section_heading(doc, "Projects & Experience", color=(10, 102, 194))
    _add_bullet_list(doc, data.get('experience',[]))
    _add_bullet_list(doc, data.get('projects',[]))

    _add_section_heading(doc, "Education", color=(10, 102, 194))
    _add_bullet_list(doc, data.get('education',[]))

    _add_section_heading(doc, "Technical Skills", color=(10, 102, 194))
    _add_bullet_list(doc, data.get('skills',[]))
    
    _add_section_heading(doc, "Certifications", color=(10, 102, 194))
    _add_bullet_list(doc, data.get('certificates', []))

    _add_section_heading(doc, "Personal Details", color=(255, 102, 0))
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")

    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(data.get('declaration', "I hereby declare that the information provided is true to the best of my knowledge."))
    return doc

def template_creative(data):
    doc = Document()
    _add_heading(doc, data.get('name',''), size=26, color=(255, 102, 0))
    doc.add_paragraph(data.get('title',''))
    doc.add_paragraph(data.get('contact',''))

    _add_section_heading(doc, "About Me", color=(255, 102, 0))
    doc.add_paragraph(data.get('summary',''))
    
    _add_section_heading(doc, "Experience Highlights", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('experience',[]))

    _add_section_heading(doc, "Education Path", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('education',[]))

    _add_section_heading(doc, "Skillset", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('skills',[]))

    _add_section_heading(doc, "Projects", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('projects', []))

    _add_section_heading(doc, "Certifications", color=(255, 102, 0))
    _add_bullet_list(doc, data.get('certificates', []))

    _add_section_heading(doc, "Personal Details", color=(255, 102, 0))
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")

    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(data.get('declaration', "I hereby declare that the information provided is true to the best of my knowledge."))
    return doc

def template_infographic(data):
    doc = Document()
    _add_heading(doc, data.get('name',''), size=24, color=(76, 175, 80))
    doc.add_paragraph(data.get('title',''))
    doc.add_paragraph(data.get('contact',''))

    _add_section_heading(doc, "Snapshot", color=(76, 175, 80))
    doc.add_paragraph(data.get('summary',''))
    
    _add_section_heading(doc, "Key Experiences", color=(76, 175, 80))
    _add_bullet_list(doc, data.get('experience',[]))

    _add_section_heading(doc, "Learning", color=(76, 175, 80))
    _add_bullet_list(doc, data.get('education',[]))

    _add_section_heading(doc, "Proficiencies", color=(76, 175, 80))
    _add_bullet_list(doc, data.get('skills',[]))

    _add_section_heading(doc, "Projects", color=(76, 175, 80))
    _add_bullet_list(doc, data.get('projects', []))
    
    _add_section_heading(doc, "Certifications", color=(76, 175, 80))
    _add_bullet_list(doc, data.get('certificates', []))

    _add_section_heading(doc, "Personal Info", color=(255, 102, 0))
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")

    _add_section_heading(doc, "Declaration", color=(255, 102, 0))
    doc.add_paragraph(data.get('declaration', "I hereby declare that the information provided is true to the best of my knowledge."))
    return doc

def template_bordered(data):
    doc = Document()
    doc.add_paragraph("_" * 100)
    
    _add_heading(doc, data.get('name',''), size=20, color=(0, 0, 0))
    doc.add_paragraph(data.get('title',''))
    doc.add_paragraph(data.get('contact',''))

    _add_section_heading(doc, "Summary")
    doc.add_paragraph(data.get('summary',''))
    _add_section_heading(doc, "Experience")
    _add_bullet_list(doc, data.get('experience',[]))
    _add_section_heading(doc, "Education")
    _add_bullet_list(doc, data.get('education',[]))
    _add_section_heading(doc, "Skills")
    _add_bullet_list(doc, data.get('skills',[]))
    _add_section_heading(doc, "Projects")
    _add_bullet_list(doc, data.get('projects',[]))
    _add_section_heading(doc, "Certifications & Achievements")
    _add_bullet_list(doc, data.get('certificates',[]))
    _add_section_heading(doc, "Hobbies & Interests")
    doc.add_paragraph(data.get('hobbies',''))
    _add_section_heading(doc, "Personal Details")
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")
    _add_section_heading(doc, "Declaration")
    doc.add_paragraph(data.get('declaration',''))
    
    doc.add_paragraph("_" * 100)
    return doc

def template_side_panel(data):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)
    
    _add_heading(doc, data.get('name', ''), size=22, color=(33, 33, 33))
    para = doc.add_paragraph(f"{data.get('title', '')}\n{data.get('contact', '')}")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(2.3)
    table.columns[1].width = Inches(6.0)

    left = table.cell(0, 0)
    p = left.add_paragraph(); run = p.add_run("SKILLS"); run.bold = True
    _add_bullet_list(left, data.get('skills', []))
    p = left.add_paragraph(); run = p.add_run("HOBBIES & INTERESTS"); run.bold = True
    left.add_paragraph(data.get('hobbies', ''))
    p = left.add_paragraph(); run = p.add_run("PERSONAL DETAILS"); run.bold = True
    for key, value in data.get('personal_details', {}).items():
        left.add_paragraph(f"{key}: {value}")

    right = table.cell(0, 1)
    p = right.add_paragraph(); run = p.add_run("PROFILE SUMMARY"); run.bold = True
    right.add_paragraph(data.get('summary', ''))
    p = right.add_paragraph(); run = p.add_run("EXPERIENCE"); run.bold = True
    _add_bullet_list(right, data.get('experience', []))
    p = right.add_paragraph(); run = p.add_run("EDUCATION"); run.bold = True
    _add_bullet_list(right, data.get('education', []))
    p = right.add_paragraph(); run = p.add_run("PROJECTS"); run.bold = True
    _add_bullet_list(right, data.get('projects', []))
    p = right.add_paragraph(); run = p.add_run("CERTIFICATIONS"); run.bold = True
    _add_bullet_list(right, data.get('certificates', []))
    p = right.add_paragraph(); run = p.add_run("DECLARATION"); run.bold = True
    right.add_paragraph(data.get('declaration', ''))
    
    return doc

def template_academic(data):
    doc = Document()
    _add_heading(doc, data.get('name',''), size=18, color=(54, 69, 79))
    doc.add_paragraph(data.get('title',''))
    doc.add_paragraph(data.get('contact',''))
    
    _add_section_heading(doc, "Research Profile", color=(54, 69, 79))
    doc.add_paragraph(data.get('summary',''))
    _add_section_heading(doc, "Teaching / Research Experience", color=(54, 69, 79))
    _add_bullet_list(doc, data.get('experience',[]))
    _add_section_heading(doc, "Education", color=(54, 69, 79))
    _add_bullet_list(doc, data.get('education',[]))
    _add_section_heading(doc, "Publications & Skills", color=(54, 69, 79))
    _add_bullet_list(doc, data.get('skills',[]))
    _add_section_heading(doc, "Projects")
    _add_bullet_list(doc, data.get('projects',[]))
    _add_section_heading(doc, "Certifications & Achievements")
    _add_bullet_list(doc, data.get('certificates',[]))
    _add_section_heading(doc, "Hobbies & Interests")
    doc.add_paragraph(data.get('hobbies',''))
    _add_section_heading(doc, "Personal Details")
    for key, value in data.get('personal_details', {}).items():
        doc.add_paragraph(f"{key}: {value}")
    _add_section_heading(doc, "Declaration")
    doc.add_paragraph(data.get('declaration',''))
    return doc

def template_executive(data):
    doc = Document()
    _add_heading(doc, data.get('name', ''), size=20, color=(0,0,0))
    doc.add_paragraph(data.get('title', ''))
    doc.add_paragraph(data.get('contact', ''))
    
    _add_section_heading(doc, "Executive Summary", color=(80,80,80))
    doc.add_paragraph(data.get('summary', ''))
    _add_section_heading(doc, "Professional Experience", color=(80,80,80))
    _add_bullet_list(doc, data.get('experience', []))
    _add_section_heading(doc, "Education", color=(80,80,80))
    _add_bullet_list(doc, data.get('education', []))
    _add_section_heading(doc, "Key Skills", color=(80,80,80))
    _add_bullet_list(doc, data.get('skills', []))
    
    return doc

def template_simple_ats(data):
    doc = Document()
    doc.add_heading(data.get('name', 'Your Name'), level=1)
    doc.add_paragraph(f"Title: {data.get('title', '')}")
    doc.add_paragraph(f"Contact: {data.get('contact', '')}")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(data.get('summary', ''))
    doc.add_heading("Experience", level=2)
    for exp in data.get('experience', []):
        doc.add_paragraph(exp, style="List Bullet")
    doc.add_heading("Education", level=2)
    for edu in data.get('education', []):
        doc.add_paragraph(edu, style="List Bullet")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(data.get('skills', [])))
    doc.add_heading("Projects", level=2)
    for proj in data.get('projects', []):
        doc.add_paragraph(proj, style="List Bullet")
    return doc


# Template mapping
TEMPLATE_MAP = {
    "Minimal": template_minimal,
    "Corporate": template_corporate,
    "Tech Modern": template_tech_modern,
    "Creative": template_creative,
    "Infographic Style": template_infographic,
    "Simple Bordered": template_bordered,
    "Side Panel": template_side_panel,
    "Academic / Research": template_academic,
    "Executive": template_executive,
    "Simple ATS-Friendly": template_simple_ats,
}

# ---------------- Utility helpers ----------------
def save_doc_to_link(doc, filename="resume.docx"):
    bio = io.BytesIO()
    doc.save(bio)
    b64 = base64.b64encode(bio.getvalue()).decode()
    link = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="card-cta" style="text-decoration:none;display:inline-block;margin-top:12px;">📥 Download Resume</a>'
    return link

def pretty_profile_html(profile):
    name = profile.get("name", "N/A")
    headline = profile.get("headline", "")
    location = profile.get("location", "")
    skills = ", ".join(profile.get("skills", []))
    exps = profile.get("experiences", [])
    edus = profile.get("educations", [])

    exp_html = "".join([f"<li><b>{exp.get('title','')}</b> at {exp.get('company','')} ({exp.get('date_range','')})<br><small>{exp.get('summary', '')}</small></li>" for exp in exps]) or "<li>Not found</li>"
    edu_html = "".join([f"<li><b>{edu.get('degree','')}</b>, {edu.get('field','')} at {edu.get('school','')} ({edu.get('date_range','')})</li>" for edu in edus]) or "<li>Not found</li>"

    return f"""
    <div style="border:1px solid #ddd; border-radius:10px; padding:20px; margin-top:15px; background-color: #fafafa;">
        <h3>{name}</h3>
        <p><b>{headline}</b><br><i>{location}</i></p>
        <hr>
        <h4>Skills</h4><p>{skills or 'Not found'}</p>
        <h4>Experience</h4><ul>{exp_html}</ul>
        <h4>Education</h4><ul>{edu_html}</ul>
    </div>
    """

def create_driver(headless=True, implicit_wait=10):
    opts = Options()
    if headless: opts.add_argument("--headless=new")
    else: opts.add_argument("--start-maximized")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--disable-extensions")
    opts.add_argument("--disable-blink-features=AutomationControlled")
    opts.add_argument("--window-size=1400,1000")
    opts.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36")
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=opts)

def fetch_public_html(url, timeout=10):
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36"}
    try:
        r = requests.get(url, headers=headers, timeout=timeout)
        if r.status_code == 200 and len(r.text) > 3000 and "sign in" not in r.text.lower():
            return r.text
    except Exception as e:
        logging.info("requests fetch failed: %s", e)
    return None

def fetch_profile_html_selenium(url, headless=True):
    """
    More robust Selenium fetcher that waits for the main profile element to load.
    """
    driver = None
    try:
        driver = create_driver(headless=headless)
        driver.get(url)
        # Wait for the main profile container to be visible before parsing
        WebDriverWait(driver, 15).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "main.scaffold-layout__main"))
        )
        html = driver.page_source
        return html
    except Exception as e:
        logging.error(f"Selenium fetch failed: {e}")
        if driver:
            return driver.page_source
        return None
    finally:
        if driver:
            driver.quit()


def safe_text(sel):
    return sel.get_text(strip=True) if sel else ""

def parse_linkedin_profile(html, profile_url):
    """
    Parses LinkedIn profile HTML, updated with more robust 2024 selectors.
    Handles cases where sections might be missing.
    """
    soup = BeautifulSoup(html, "html.parser")
    
    main_content = soup.select_one("main.scaffold-layout__main")
    if not main_content:
        main_content = soup

    # --- Basic Info ---
    top_card = main_content.find("section", class_=lambda x: x and "pv-top-card" in x)
    name = ""
    headline = ""
    location = ""
    if top_card:
        name = safe_text(top_card.select_one("h1.text-heading-xlarge"))
        headline = safe_text(top_card.select_one("div.text-body-medium"))
        location = safe_text(top_card.select_one("span.text-body-small.inline"))

    # --- Experience Section ---
    experiences = []
    exp_section = main_content.find("section", id="experience")
    if exp_section:
        for item in exp_section.select("li.pvs-list__paged-list-item, div.pv-entity__position-group-pager"):
            title = safe_text(item.select_one("h3, .t-bold span[aria-hidden='true']"))
            company_line = item.select_one(".t-normal span[aria-hidden='true']")
            company = ""
            if company_line:
                company = company_line.get_text(strip=True).split('·')[0].strip()
            date_range = safe_text(item.select_one("h4.pv-entity__date-range span:nth-child(2), .t-black--light span[aria-hidden='true']"))
            summary = safe_text(item.select_one("p.pv-entity__description, .pvs-entity__description"))
            if title and company:
                experiences.append({"title": title, "company": company, "date_range": date_range, "summary": summary})

    # --- Education Section ---
    educations = []
    edu_section = main_content.find("section", id="education")
    if edu_section:
        for item in edu_section.select("li.pvs-list__paged-list-item"):
            school = safe_text(item.select_one(".t-bold span[aria-hidden='true']"))
            degree_line = item.select(".t-normal span[aria-hidden='true']")
            degree, field = "", ""
            if len(degree_line) > 0: degree = safe_text(degree_line[0])
            if len(degree_line) > 1: field = safe_text(degree_line[1])
            date_range = safe_text(item.select_one(".t-black--light span[aria-hidden='true']"))
            if school:
                educations.append({"school": school, "degree": degree, "field": field, "date_range": date_range})

    # --- Skills Section ---
    skills = []
    skills_section = main_content.find("section", id="skills")
    if skills_section:
        for skill_item in skills_section.select(".pvs-skill-category-entity__name-text, .skill-pill, .pv-skill-entity__skill-name"):
            skills.append(safe_text(skill_item))
    
    return {
        "name": name, "headline": headline, "location": location,
        "experiences": experiences, "educations": educations,
        "skills": list(set(skills)), "profile_url": profile_url
    }

def build_template_data_from_profile(profile):
    return {
        "name": profile.get("name", ""),
        "title": profile.get("headline", ""),
        "contact": profile.get("location", ""),
        "summary": profile.get("headline", ""),
        "experience": [f"{exp.get('title', '')} at {exp.get('company', '')} ({exp.get('date_range', '')})\n{exp.get('summary', '')}" for exp in profile.get("experiences", [])],
        "education": [f"{edu.get('degree', '')}, {edu.get('field','')} at {edu.get('school', '')} ({edu.get('date_range', '')})" for edu in profile.get("educations", [])],
        "skills": profile.get("skills", []),
        "projects": [], "certificates": [], "personal_details": {}, "declaration": "", "hobbies": ""
    }

# ---------------- UI styling and Layout ----------------
st.markdown("""
<style>
/* Main app styling */
.stApp {
    background: linear-gradient(120deg, #a1c4fd 0%, #c2e9fb 50%, #fbc2eb 100%);
    background-size: 200% 200%;
    animation: gradientMove 10s ease-in-out infinite;
    color: #07203a;
}
@keyframes gradientMove { 0%{background-position:0% 50%} 50%{background-position:100% 50%} 100%{background-position:0% 50%} }

/* Hero section for the header */
.hero { 
    text-align:center; 
    padding: 1rem 1rem; 
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 20px;
}
.hero-text {
    max-width: 600px;
}
.hero-title { 
    font-size:2.8rem; 
    font-weight:900; 
    color:#4b006e; 
    text-shadow: 0 2px 10px #a1c4fd88; 
    margin-bottom: 0.5rem;
}
.hero-sub { 
    color:#274060; 
    font-size:1.1rem; 
    margin-bottom:1rem; 
}

/* Main content cards grid */
.cards-grid {
    display: grid;
    grid-template-columns: repeat(2, 1fr);
    gap: 2rem;
    max-width: 900px; 
    margin: 2rem auto;
    padding: 0 1rem;
}

/* Card styling for navigation and content pages */
.big-card {
    background: rgba(255,255,255,0.7);
    backdrop-filter: blur(10px);
    border-radius: 22px;
    padding: 2rem;
    box-shadow: 0 10px 30px rgba(0,0,0,0.1);
    transition: all .3s ease;
    border: 1px solid rgba(255,255,255,0.2);
    display: flex;
    flex-direction: column;
}
.big-card:hover {
    transform: translateY(-10px) scale(1.02);
    box-shadow: 0 20px 40px rgba(0,0,0,0.15);
}
.big-card .icon {
    font-size: 2.5rem;
    display: block;
    margin-bottom: 1rem;
    line-height: 1;
}
.big-card .title {
    font-size: 1.5rem;
    font-weight: 800;
    color: #4b006e;
}
.big-card .desc {
    color: #2d3f57;
    margin-top: 0.5rem;
    flex-grow: 1; /* Pushes button down */
}
.card-cta {
    margin-top: 1.5rem;
    display: inline-block;
    padding: 0.6rem 1.2rem;
    border-radius: 12px;
    background: linear-gradient(90deg,#a1c4fd,#fbc2eb);
    color: white !important; /* Important to override link styles */
    text-decoration: none;
    font-weight: 700;
    box-shadow: 0 4px 15px #a1c4fd66;
    transition: all .3s;
    text-align: center;
}
.card-cta:hover {
    transform: scale(1.1);
    box-shadow: 0 6px 20px #fbc2eb88;
}

/* Score card for ATS feature */
.score-card {
    padding: 1.5rem;
    border-radius: 14px;
    text-align: center;
    background: rgba(255, 255, 255, 0.8);
    color: #4b006e;
    margin-top: 1rem;
}

/* Responsive design for smaller screens */
@media (max-width: 768px) {
    .cards-grid {
        grid-template-columns: 1fr;
    }
    .hero {
        flex-direction: column;
    }
    .hero-title {
        font-size: 2.2rem;
    }
}
</style>
""", unsafe_allow_html=True)


# ---------------- Header / Hero ----------------
lottie_resume = load_lottieurl("https://assets5.lottiefiles.com/packages/lf20_fcfjwiyb.json")
with st.container():
    st.markdown('<div class="hero">', unsafe_allow_html=True)
    c1, c2 = st.columns([1, 2])
    with c1:
        if lottie_resume:
            st_lottie(lottie_resume, speed=1, height=200, key="initial")
    with c2:
        st.markdown('<div class="hero-text"><div class="hero-title">ResuMind - Smart Resume Builder & Scorer </div><div class="hero-sub">From Crafting to Scoring — Your Entire Resume Journey in One App </div></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
st.write("---")

# ---------------- App Navigation and Pages ----------------
query_params = st.experimental_get_query_params()
if "nav" in query_params:
    st.session_state["page"] = query_params["nav"][0]
elif "page" not in st.session_state:
    st.session_state["page"] = "home"

# --- HOME PAGE ---
if st.session_state["page"] == "home":
    try:
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
    except:
        local_ip = "localhost"
    app_url = f"http://{local_ip}:{st.get_option('server.port') or 8501}"

    st.markdown(f"""
    <div class="cards-grid">
      <div class="big-card"><div class="icon">✍️</div><div class="title">Create Resume</div><div class="desc">Start from scratch, fill in your details, and pick a professional template.</div><a class="card-cta" href="{app_url}?nav=create" target="_self">Start Creating</a></div>
      <div class="big-card"><div class="icon">🔗</div><div class="title">Import from LinkedIn</div><div class="desc">Automatically generate a resume by pasting a public LinkedIn profile URL.</div><a class="card-cta" href="{app_url}?nav=linkedin" target="_self">Import Profile</a></div>
      <div class="big-card"><div class="icon">📤</div><div class="title">Upload & Edit</div><div class="desc">Upload a DOCX resume, make quick edits, and re-download.</div><a class="card-cta" href="{app_url}?nav=upload" target="_self">Upload & Edit</a></div>
      <div class="big-card"><div class="icon">📊</div><div class="title">Check ATS Score</div><div class="desc">Analyze your resume against a job description for keyword optimization.</div><a class="card-cta" href="{app_url}?nav=ats" target="_self">Analyze Now</a></div>
    </div>
    """, unsafe_allow_html=True)

# --- CREATE PAGE ---
if st.session_state["page"] == "create":
    st.markdown('<div class="big-card" style="max-width:1000px; margin:auto;"><h3 class="title">✍️ Create Your Resume</h3><p class="desc">Fill out the fields below and choose a template to generate your document.</p>', unsafe_allow_html=True)
    with st.form("create_form"):
        c1, c2 = st.columns(2)
        name = c1.text_input("Full Name", st.session_state.get("name", ""))
        title = c2.text_input("Job Title / Headline", st.session_state.get("title", ""))
        contact = st.text_area("Contact Info (Email, Phone, Location)", st.session_state.get("contact", ""), height=100)
        summary = st.text_area("Profile Summary / Career Objective", st.session_state.get("summary", ""), height=150)
        experience = st.text_area("Experience (one entry per line)", st.session_state.get("experience", ""), height=200)
        education = st.text_area("Education (one entry per line)", st.session_state.get("education", ""), height=150)
        skills = st.text_area("Skills (one per line)", st.session_state.get("skills", ""), height=150)
        projects = st.text_area("Projects (one per line)", st.session_state.get("projects", ""), height=150)
        certificates = st.text_area("Certificates/Achievements (one per line)", st.session_state.get("certificates", ""), height=100)
        personal_details = st.text_area("Personal Details (format: 'Key: Value' on each line)", st.session_state.get("personal_details", ""), height=100)
        hobbies = st.text_area("Hobbies & Interests", st.session_state.get("hobbies", ""), height=50)
        declaration = st.text_area("Declaration", st.session_state.get("declaration", "I hereby declare that the information provided is true to the best of my knowledge."), height=100)

        template_choice = st.selectbox("Choose a Resume Template", list(TEMPLATE_MAP.keys()))
        submitted = st.form_submit_button("Generate Resume Document")

    if submitted:
        data = {
            "name": name, "title": title, "contact": contact, "summary": summary, "hobbies": hobbies, "declaration": declaration,
            "experience": [e.strip() for e in experience.splitlines() if e.strip()],
            "education": [e.strip() for e in education.splitlines() if e.strip()],
            "skills": [s.strip() for s in skills.splitlines() if s.strip()],
            "projects": [p.strip() for p in projects.splitlines() if p.strip()],
            "certificates": [c.strip() for c in certificates.splitlines() if c.strip()],
            # CORRECTED: Safely parse personal details
            "personal_details": {
                k.strip(): v.strip() for k, v in (
                    line.split(":", 1) for line in personal_details.splitlines() if ":" in line
                )
            },
        }
        try:
            doc = TEMPLATE_MAP[template_choice](data)
            st.markdown(save_doc_to_link(doc, filename=f"{name.replace(' ','_')}_Resume.docx"), unsafe_allow_html=True)
            st.success("Your resume is ready for download!")
            st.balloons()
        except Exception as e:
            st.error(f"An error occurred during template generation: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

# --- LINKEDIN IMPORT PAGE ---
if st.session_state["page"] == "linkedin":
    st.markdown('<div class="big-card" style="max-width:1000px;margin:auto;"><h3 class="title">🔗 Import from LinkedIn</h3>', unsafe_allow_html=True)
    st.info("ℹ️ Provide a **public** LinkedIn profile URL. Scraping is best-effort and may fail if LinkedIn blocks the request or changes its layout.")
    
    profile_url = st.text_input("LinkedIn Profile URL", placeholder="https://www.linkedin.com/in/your-profile-name")
    template_choice = st.selectbox("Choose a Template For Your Resume", list(TEMPLATE_MAP.keys()))
    
    if st.button("Scrape & Generate Resume"):
        if not profile_url or "linkedin.com" not in profile_url:
            st.error("Please enter a valid LinkedIn profile URL.")
        else:
            html = None
            with st.spinner("Attempting to fetch profile... This may take a moment."):
                html = fetch_public_html(profile_url)
                if not html:
                    st.info("Initial fetch failed. Retrying with browser automation (this is slower)...")
                    html = fetch_profile_html_selenium(profile_url)

            if not html or "authwall" in html:
                st.error("Failed to retrieve profile. This can happen if the profile is private, the URL is wrong, or LinkedIn is blocking requests.")
            else:
                st.success("Profile HTML retrieved. Parsing data...")
                with st.spinner("Extracting information from the profile..."):
                    profile_data = parse_linkedin_profile(html, profile_url)

                if not profile_data.get('name'):
                     st.warning("Could not parse profile data completely. The generated resume may be incomplete. This might be due to a change in LinkedIn's layout.")
                else:
                    st.success(f"Successfully parsed data for {profile_data.get('name', 'the user')}.")
                
                st.markdown("### Scraped Data Preview")
                st.markdown(pretty_profile_html(profile_data), unsafe_allow_html=True)
                
                template_data = build_template_data_from_profile(profile_data)
                doc = TEMPLATE_MAP[template_choice](template_data)
                
                st.markdown("---")
                st.markdown(save_doc_to_link(doc, filename=f"{template_data.get('name', 'resume').replace(' ','_')}_Resume.docx"), unsafe_allow_html=True)
                st.balloons()

    st.markdown('</div>', unsafe_allow_html=True)

# --- UPLOAD & EDIT PAGE ---
if st.session_state["page"] == "upload":
    st.markdown('<div class="big-card" style="max-width:1000px;margin:auto;"><h3 class="title">📤 Upload & Edit Resume</h3>', unsafe_allow_html=True)
    uploaded = st.file_uploader("Upload a DOCX file to edit", type=["docx"])

    if uploaded:
        try:
            doc = Document(uploaded)
            content = "\n".join([p.text for p in doc.paragraphs])
            
            # CORRECTED: 'edited_text' is defined here by the text_area widget
            edited_text = st.text_area("Edit resume text below", value=content, height=500)

            if st.button("Save Edited Resume"):
                new_doc = Document()
                # CORRECTED: The loop correctly uses 'edited_text' which is now in scope
                for line in edited_text.splitlines():
                    new_doc.add_paragraph(line)

                st.markdown(save_doc_to_link(new_doc, filename="Edited_Resume.docx"), unsafe_allow_html=True)
                st.success("Your edited resume is ready for download!")

        except Exception as e:
            st.error(f"Failed to load or edit the .docx file: {e}")

    st.markdown('</div>', unsafe_allow_html=True)

# --- ATS ANALYZER PAGE ---
if st.session_state["page"] == "ats":
    st.markdown('<div class="big-card" style="max-width:1000px;margin:auto;"><h3 class="title">📊 ATS Keyword Analyzer</h3>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    uploaded_resume = c1.file_uploader("Upload Your Resume", type=["pdf", "docx"])
    job_description = c2.text_area("Paste Job Description", height=300, placeholder="Paste the full job description here...")

    if st.button("Analyze Match Score"):
        if not uploaded_resume or not job_description:
            st.error("Please upload a resume and paste a job description.")
        else:
            resume_text = ""
            try:
                if uploaded_resume.name.endswith(".pdf"):
                    reader = PdfReader(uploaded_resume)
                    resume_text = "".join(page.extract_text() for page in reader.pages)
                else:
                    doc = Document(uploaded_resume)
                    resume_text = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                st.error(f"Error reading resume file: {e}")

            if resume_text:
                jd_words = set(w.lower() for w in job_description.split() if len(w) > 3 and w.isalpha())
                resume_lower = resume_text.lower()
                hits = [w for w in jd_words if w in resume_lower]
                score = round(len(hits) / max(1, len(jd_words)) * 100)
                
                st.markdown(f'<div class="score-card"><h2>ATS Match Score: {score}%</h2><p>Found {len(hits)} of {len(jd_words)} potential keywords.</p></div>', unsafe_allow_html=True)
                
                c1, c2 = st.columns(2)
                c1.success(f"**Matched Keywords:** {', '.join(hits)}")
                missed = list(jd_words - set(hits))
                c2.warning(f"**Keywords to Consider:** {', '.join(missed[:20])}") # Show top 20 missed

    st.markdown('</div>', unsafe_allow_html=True)